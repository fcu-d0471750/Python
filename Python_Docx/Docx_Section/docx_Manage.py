'''
doc管理
'''
import docx
#加入可調整的 word 單位
from docx.shared import Cm
#處理文件的直向/橫向
from docx.enum.section import WD_ORIENTATION
#加入可調整的 word 單位
from docx.enum.text import WD_TAB_ALIGNMENT
# 導入資料庫管理
from Store_DB.DB_Manage import use_store_db
# 導入亂數
import random


#=================================================================
# 宣告變數
#=================================================================

# doc預設格式
doc = docx.Document('../pages/M.docx')

# 輸出doc的數量
doc_count = 0

# 符合 require_store 的門市資料
allow_store = []

# doc的調整係數
section = doc.sections[0]

# 表格設定(rows=多的列的數量,cols=固定的行數量)
# rows=1表示不多新增列，也就是只依照records的列的數量
table = doc.add_table(rows=1, cols=11)

#=================================================================
# 宣告副程式
#=================================================================

# 預設值
def ini_doc():
    # 使用外部的全欲變數
    global allow_store,doc,table

    # 清空符合 require_store 的門市資料
    allow_store = []
    # doc恢復成預設
    doc = docx.Document('../pages/M.docx')
    # 表格設定(rows=多的列的數量,cols=固定的行數量)
    # rows=1表示不多新增列，也就是只依照records的列的數量
    table = doc.add_table(rows=1, cols=11)

# 調整 doc 係數
def set_section():
    # 使用外部的全欲變數
    global section

    # doc的調整
    section = doc.sections[0]
    # 調整橫式紙張
    section.orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = Cm(29.7), Cm(21)
    section.page_width = new_width
    section.page_height = new_height
    # 調整文件左右上下邊界至 1.27 cm
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

# 執行 doc
def do_doc(require_area,require_store):
    # 輸出doc的數量
    doc_count = len(require_store)

    # 共執行幾次
    for count in range(0, doc_count):
        # 恢復成預設
        ini_doc()
        # 調整 doc 係數
        set_section()


        # 找出符合 require_store 的門市資料
        for bbb in require_store[count]:
            for j in range(0, len(use_store_db(require_area))):
                if (bbb == use_store_db(require_area)[j][2]):
                    allow_store.append(use_store_db(require_area)[j])

        # 寫入doc
        for Zone, cls, name, date, brand, number, radiation, power, Components_01, Components_02, question_01 in allow_store:
            row_cells = table.add_row().cells
            row_cells[0].text = Zone
            row_cells[1].text = cls
            row_cells[2].text = name
            row_cells[3].text = date
            row_cells[4].text = brand
            row_cells[5].text = number
            row_cells[6].text = ('%.2f' % random.uniform(0.11, 0.29))  # 保留小數點第二位，並保留0
            row_cells[7].text = str(random.randrange(1250, 1350))
            row_cells[8].text = Components_01
            row_cells[9].text = Components_02
            row_cells[10].text = question_01

        # 輸出的doc的名稱
        doc_name = 'AI 技術可以讓隱藏於暗處的物品現形' + str(count) + '.docx'
        # 輸出 docx
        doc.save(doc_name)














